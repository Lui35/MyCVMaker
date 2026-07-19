from io import BytesIO
from pathlib import Path

import pymupdf
from docx import Document


MAX_DOCUMENT_BYTES = 8 * 1024 * 1024
MAX_EXTRACTED_CHARACTERS = 60_000
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


class DocumentExtractionError(ValueError):
    pass


def extract_document_text(filename: str, content: bytes) -> tuple[str, list[str]]:
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise DocumentExtractionError("Use a PDF, DOCX, or TXT file.")
    if not content:
        raise DocumentExtractionError("The uploaded file is empty.")
    if len(content) > MAX_DOCUMENT_BYTES:
        raise DocumentExtractionError("The uploaded file must be smaller than 8 MB.")

    warnings: list[str] = []
    if extension == ".pdf":
        text = _extract_pdf(content)
        if len(text.strip()) < 120:
            raise DocumentExtractionError(
                "This PDF appears to be scanned or image-only. OCR support is the next import step; upload a text-based PDF, DOCX, or TXT file for now."
            )
    elif extension == ".docx":
        text = _extract_docx(content)
    else:
        text = content.decode("utf-8", errors="replace")

    normalized = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if len(normalized) < 80:
        raise DocumentExtractionError("Not enough readable CV text was found in the uploaded file.")
    if len(normalized) > MAX_EXTRACTED_CHARACTERS:
        normalized = normalized[:MAX_EXTRACTED_CHARACTERS]
        warnings.append("Only the first 60,000 characters were imported.")
    return normalized, warnings


def _extract_pdf(content: bytes) -> str:
    try:
        with pymupdf.open(stream=content, filetype="pdf") as document:
            return "\n".join(page.get_text("text", sort=True) for page in document)
    except Exception as exc:
        raise DocumentExtractionError("The PDF could not be read.") from exc


def _extract_docx(content: bytes) -> str:
    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise DocumentExtractionError("The DOCX file could not be read.") from exc

    blocks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            blocks.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
    return "\n".join(blocks)
